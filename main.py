# from docx import Document
# import paragraphs
#
# info = ['Index', 'Name', 'Surname']
# data = [
#     [1, 'Mannopboy', 'Mukinboyev'],
#     [2, 'Elyor', 'Xamidulayev'],
#     [3, 'Xudoyor', 'Tursunov'],
#     [4, 'Shohruh', 'Shermetov'],
# ]
#
# document = Document()
# document.add_heading('He MMR', 0)
# document.add_heading('He MMR', 1)
# document.add_heading('He MMR', 2)
# document.add_heading('He MMR', 3)
# document.add_heading('He MMR', 4)
# document.add_heading('He MMR', 5)
#
# p = document.add_paragraph('New text!!!')
#
# p.add_run('Text 1').bold = True
# p.add_run('Text 1').italic = True
#
# table = document.add_table(rows=1, cols=3)
#
# for item in range(3):
#     table.rows[0].cells[item].text = info[item]
#
# for index, name, surname in data:
#     calls = table.add_row().cells
#     calls[0].text = str(index)
#     calls[1].text = name
#     calls[2].text = surname
#
# document.add_page_break()
#
# document.add_paragraph('New page')
#
# document.save('my.docx')

from docx.api import Document
document = Document('my.docx')
